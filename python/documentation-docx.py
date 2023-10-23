from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml import OxmlElement

# CRIA UM NOVO DOCUMENTO
doc = Document()

name = "Gustavo de Souza Fonseca"
access = ["MV", "DeskManager", "Email", "CWM", "Rede", "Interact"]
email = "gustavo.fonseca@hnipo.org.br"
username = "HC51132013"
password = "Hnb@511*0109"
acessos_por_email = ""
acessos_por_login = access


def formatar_acessos_por_email(acessos: list):
    if isinstance(acessos, list):
        if "DeskManager" in acessos and "Email" not in acessos:
            return "DeskManager"
        elif "Email" in acessos and "DeskManager" not in acessos:
            return "Email"
        elif "DeskManager" in acessos and "Email" in acessos:
            return "DeskManager, Email"
    return "Acesso não especificado"  # or any default value you prefer



def formatar_acessos_por_login(acessos):
    if "DeskManager" in acessos:
        acessos_por_login.remove("DeskManager")
    if "Email" in acessos:
        acessos_por_login.remove("Email")



# CRIANDO FUNÇÃO DE UM NOVO TEXTO FORMATADO COM UM PADRÃO CASO NÃO SEJA PASSADO INFORMAÇÕES ABAIXO
def texto_formatado(
    document,
    text,
    font_name="Calibri",
    bold=False,
    underline=False,
    alignment=None,
    font_size=12,
):
    paragraph = document.add_paragraph(text)
    for run in paragraph.runs:
        run.bold = bold
        run.font.underline = underline
        if font_size:
            run.font.size = Pt(font_size)
        if font_name:
            run.font.name = font_name

    if alignment:
        paragraph.alignment = alignment

    return paragraph


# CRIADO UM HEADER PADRÃO DO WORD
header = doc.sections[0].header


# ADICIONANDO IMAGEM NO HEADER
image_path = './assets/logotipo.jpg'
header_paragraph = header.paragraphs[0]
run = header_paragraph.add_run()
run.add_picture(image_path, width=Inches(3.5))


# CENTRALIZANDO TODO O CONTEÚDO DO HEADER
for paragraph in header.paragraphs:
    paragraph.alignment = 1

doc.add_paragraph()

texto_formatado(
    doc,
    "Acessos Corporativos",
    bold=True,
    underline=True,
    alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
    font_size=16,
)

texto_formatado(
    doc, "               Prezado (a),", font_size=11, font_name="Calibri (Corpo)"
)
texto_formatado(doc, f"                    {name},", bold=True, font_size=14)
doc.add_paragraph()

texto_formatado(
    doc,
    "     Em nome do departamento de Tecnologia da Informação, gostaríamos de lhe desejar boas-vindas e muito sucesso em nossa instituição.",
)
doc.add_paragraph()
texto_formatado(
    doc,
    "     Abaixo compartilhamos os dados de acesso de algumas ferramentas que estão sob a gestão da T.I.:",
)
doc.add_paragraph()

# with open(file, encoding="utf-8") as csvfile:
#     reader = csv.DictReader(csvfile, delimiter=";")

# Crie uma nova lista com marcadores e aninhe os detalhes e usuários
doc.add_paragraph("Sistemas:", style="ListBullet")
doc.add_paragraph(f"{formatar_acessos_por_login(access)}", style="ListBullet2")
doc.add_paragraph(f"{username}", style="ListBullet3")
doc.add_paragraph(f"{formatar_acessos_por_email(access)}:", style="ListBullet2")
doc.add_paragraph(f"{email}", style="ListBullet3")
doc.add_paragraph(f"Senha para todos usuários:", style="ListBullet2")
doc.add_paragraph(f"{password}", style="ListBullet3")
doc.add_paragraph()
doc.add_paragraph(
    "Obs.: Após cada um dos respectivos acessos, será solicitado a alteração da senha."
)
doc.add_paragraph()
doc.add_paragraph(
    "     Caso tenha algum problema com os referidos acessos, ou qualquer outra ferramenta de trabalho administrada pela T.I., disponibilizamos o sistema DeskManager para registro de requisições e inicidentes, e também o telefone (11) 2633-2818 que estará disponível 24 horas para que possamos lhe ajudar."
)

# Crie um rodapé em todas as seções do documento
for section in doc.sections:
    footer = section.footer

    # Adicione um parágrafo ao rodapé
    footer_paragraph = (
        footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    )

    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adicione o texto "TECNOLOGIA DA INFORMAÇÃO" em negrito
    run = footer_paragraph.add_run("TECNOLOGIA DA INFORMAÇÃO")
    run.bold = True
    run.font.size = Pt(10)


doc.save(f"../ACESSOS/{name}.docx")