from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml import OxmlElement
import sys
import os

# name = sys.argv[1]
# access = sys.argv[2]
# email = sys.argv[3]
# username = sys.argv[4]
# password = sys.argv[5]




def createDoc(name, access, email, username, password):
    print(f'''
        nome: {name}
        acesso: {access}
        email: {email}
        usuario: {username}
        senha: {password}
                        ''')

    try:

        # CRIA UM NOVO DOCUMENTO
        doc = Document()
        
        acessos_por_email = ""
        acessos_por_usuario = access


        def formatar_acessos_por_email(acessos):
            acessos = acessos.split(", ")
        
            if "DeskManager" in acessos and "Conta de e-mail" not in acessos:
                acessos = "DeskManager"
                return acessos
            elif "Conta de e-mail" in acessos and "DeskManager" not in acessos:
                acessos = "Email"
                return acessos
            elif "DeskManager" in acessos and "Conta de e-mail" in acessos:
                acessos = "DeskManager, Email"
                return acessos
            return "Acesso não especificado"

        def formatar_acessos_por_usuario(acessos):
            acessos = acessos.split(", ")

            if "DeskManager" in acessos:
                acessos.remove("DeskManager")
                return acessos
            if "Conta de e-mail" in acessos:
                acessos.remove("Conta de e-mail")
                return acessos



        # CRIANDO FUNÇÃO DE UM NOVO TEXTO FORMATADO COM UM PADRÃO CASO NÃO SEJA PASSADO INFORMAÇÕES ABAIXO
        def texto_formatado(
            document,
            text,
            font_name="Calibri",
            bold=False,
            underline=False,
            alignment=None,
            font_size=12,
        ):
            paragraph = document.add_paragraph(text)
            for run in paragraph.runs:
                run.bold = bold
                run.font.underline = underline
                if font_size:
                    run.font.size = Pt(font_size)
                if font_name:
                    run.font.name = font_name

            if alignment:
                paragraph.alignment = alignment

            return paragraph


        # CRIADO UM HEADER PADRÃO DO WORD
        header = doc.sections[0].header


        # ADICIONANDO IMAGEM NO HEADER
        image_path = os.path.join(os.path.dirname(__file__), 'assets', 'logotipo.jpg')
        header_paragraph = header.paragraphs[0]
        run = header_paragraph.add_run()
        run.add_picture(image_path, width=Inches(3.5))


        # CENTRALIZANDO TODO O CONTEÚDO DO HEADER
        for paragraph in header.paragraphs:
            paragraph.alignment = 1

        doc.add_paragraph()

        texto_formatado(
            doc,
            "Acessos Corporativos",
            bold=True,
            underline=True,
            alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
            font_size=16,
        )

        texto_formatado(
            doc, "               Prezado (a),", font_size=11, font_name="Calibri (Corpo)"
        )
        texto_formatado(doc, f"                    {name},", bold=True, font_size=14)
        doc.add_paragraph()

        texto_formatado(
            doc,
            "     Em nome do departamento de Tecnologia da Informação, gostaríamos de lhe desejar boas-vindas e muito sucesso em nossa instituição.",
        )
        doc.add_paragraph()
        texto_formatado(
            doc,
            "     Abaixo compartilhamos os dados de acesso de algumas ferramentas que estão sob a gestão da T.I.:",
        )
        doc.add_paragraph()

        # Crie uma nova lista com marcadores e aninhe os detalhes e usuários
        doc.add_paragraph("Sistemas:", style="ListBullet")

        doc.add_paragraph(f"{formatar_acessos_por_usuario(access)}", style="ListBullet2")
        doc.add_paragraph(f"{username}", style="ListBullet3")

        if email is not None and email != "undefined":
            doc.add_paragraph(f"{formatar_acessos_por_email(access)}:", style="ListBullet2")
            doc.add_paragraph(f"{email}", style="ListBullet3")

        doc.add_paragraph(f"Senha para todos usuários:", style="ListBullet2")
        doc.add_paragraph(f"{password}", style="ListBullet3")

        doc.add_paragraph()
        doc.add_paragraph(
            "Obs.: Após cada um dos respectivos acessos, será solicitado a alteração da senha."
        )
        doc.add_paragraph()
        doc.add_paragraph(
            "     Caso tenha algum problema com os referidos acessos, ou qualquer outra ferramenta de trabalho administrada pela T.I., disponibilizamos o sistema DeskManager para registro de requisições e inicidentes, e também o telefone (11) 2633-2818 que estará disponível 24 horas para que possamos lhe ajudar."
        )

        # Crie um rodapé em todas as seções do documento
        for section in doc.sections:
            footer = section.footer

            # Adicione um parágrafo ao rodapé
            footer_paragraph = (
                footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            )

            footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Adicione o texto "TECNOLOGIA DA INFORMAÇÃO" em negrito
            run = footer_paragraph.add_run("TECNOLOGIA DA INFORMAÇÃO")
            run.bold = True
            run.font.size = Pt(10)


        doc.save(f"./ACESSOS/{name}.docx")
        
    except Exception as e:
        print(f"Ocorreu um erro: {e}")
        return

